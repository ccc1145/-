from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"C:\Users\tangz\Desktop\xiuxian-simulator")
DOCS = ROOT / "docs"
FILES = [
    "作业一-设计阶段.md",
    "作业二-开发阶段.md",
    "作业三-测试收尾.md",
]

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
GRAY = RGBColor(95, 99, 104)


def set_font(run, name="Calibri", size=None, bold=None, color=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "微软雅黑")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def add_page_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_font(run, size=9, color=GRAY)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)
    run = paragraph.add_run(" 页")
    set_font(run, size=9, color=GRAY)


def setup_styles(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for list_name in ["List Bullet", "List Number"]:
        style = styles[list_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167


def add_header_footer(doc, short_title):
    section = doc.sections[0]
    hp = section.header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.space_after = Pt(0)
    r = hp.add_run("《修仙模拟器》项目作业")
    set_font(r, size=9, bold=True, color=GRAY)
    r = hp.add_run(f"  |  {short_title}")
    set_font(r, size=9, color=GRAY)
    add_page_field(section.footer.paragraphs[0])


def add_title_block(doc, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(title)
    set_font(r, size=23, bold=True, color=RGBColor(0, 0, 0))

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(16)
    r = p.add_run("修仙模拟器 · 项目流程材料补交任务")
    set_font(r, size=12, color=GRAY)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(14)
    r = p.add_run("提交要求：以真实项目成果为依据整理，不重复开发，不虚构过程。")
    set_font(r, size=10.5, bold=True, color=DARK_BLUE)


def add_inline_runs(paragraph, text):
    parts = re.split(r"(`[^`]+`|\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_font(run, name="Consolas", size=9.5, color=DARK_BLUE)
        elif part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_font(run, bold=True)
        else:
            run = paragraph.add_run(part)
            set_font(run)


def build(md_path):
    lines = md_path.read_text(encoding="utf-8").splitlines()
    title = lines[0].removeprefix("# ").strip()
    doc = Document()
    setup_styles(doc)
    add_header_footer(doc, title.split("——")[0])
    add_title_block(doc, title)

    in_code = False
    for raw in lines[1:]:
        line = raw.rstrip()
        if line.startswith("```"):
            in_code = not in_code
            continue
        if not line:
            continue
        if in_code:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.space_after = Pt(10)
            r = p.add_run(line.strip())
            set_font(r, name="Consolas", size=9.5, color=DARK_BLUE)
            continue
        if line.startswith("## "):
            doc.add_paragraph(line[3:].strip(), style="Heading 1")
            continue
        if line.startswith("### "):
            doc.add_paragraph(line[4:].strip(), style="Heading 2")
            continue
        m = re.match(r"^\s*\d+\.\s+(.*)$", line)
        if m:
            p = doc.add_paragraph(style="List Number")
            add_inline_runs(p, m.group(1))
            continue
        m = re.match(r"^\s*-\s+(.*)$", line)
        if m:
            p = doc.add_paragraph(style="List Bullet")
            add_inline_runs(p, m.group(1))
            continue
        p = doc.add_paragraph()
        add_inline_runs(p, line.strip())

    out = md_path.with_suffix(".docx")
    doc.core_properties.title = title
    doc.core_properties.subject = "修仙模拟器项目阶段作业"
    doc.core_properties.author = "修仙模拟器项目组"
    doc.core_properties.keywords = "修仙模拟器, 项目作业, 流程记录"
    doc.save(out)
    return out


if __name__ == "__main__":
    for name in FILES:
        print(build(DOCS / name))
